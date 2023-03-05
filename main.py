from flask import Flask, render_template, request, redirect, url_for, send_file
from PyPDF2 import PdfReader
from docx.shared import Inches
from docx import Document
from io import BytesIO
from PIL import Image
from pdf2image import convert_from_bytes
import pytesseract
pytesseract.pytesseract.tesseract_cmd = 'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'


app = Flask(__name__)


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/convert', methods=['POST'])
def convert():
    # Verifica se o arquivo foi enviado corretamente
    if 'file' not in request.files:
        return redirect(url_for('index'))

    # Lê o arquivo enviado pelo formulário
    file = request.files['file']
    if file.filename == '':
        return redirect(url_for('index'))

    # Verifica se o arquivo é um arquivo pdf válido
    if not file.filename.endswith('.pdf'):
        return render_template('index.html', error='Por favor, selecione um arquivo .pdf')

    # Lê o conteúdo do arquivo
    pdf_bytes = file.read()

    # Verifica se o conteúdo do arquivo é válido
    if not pdf_bytes:
        return render_template('index.html', error='O arquivo está vazio')

    # Converte o arquivo pdf em texto
    reader = PdfReader(BytesIO(pdf_bytes))
    text = '\n'.join([page.extract_text() for page in reader.pages])

    # Converte o arquivo pdf em imagens
    images = convert_from_bytes(pdf_bytes)

    # Extrai o texto de cada imagem e salva em uma lista
    img_text = []
    for img in images:
        img_bytes = BytesIO()
        img.save(img_bytes, format='JPEG')
        # Adicione essa linha para redefinir a posição do ponteiro
        img_bytes.seek(0)
        img_text.append(pytesseract.image_to_string(Image.open(img_bytes)))

    # Cria um novo documento do Word
    doc = Document()

    # Adiciona o texto extraído do pdf ao documento
    doc.add_paragraph(text)

    # Adiciona as imagens ao documento
    for i, img in enumerate(images):
        img_bytes = BytesIO()
        img.save(img_bytes, format='JPEG')
        doc.add_picture(img_bytes, width=Inches(6))

        # Adiciona o texto extraído da imagem ao documento
        doc.add_paragraph(img_text[i])

    # Salva o documento em um arquivo docx
    output_filename = file.filename.replace('.pdf', '.docx')
    doc.save(output_filename)

    # Define o cabeçalho Content-Disposition para fazer o download do arquivo convertido
    mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    return send_file(output_filename, as_attachment=True, mimetype=mimetype)


if __name__ == '__main__':
    app.run(debug=True)
